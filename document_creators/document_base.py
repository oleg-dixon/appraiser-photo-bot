"""Базовые функции для создания документов Word."""

import io
import logging
from typing import List, Dict, Any, Tuple, Optional

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .constants import SIZE_OPTIONS, A4_WIDTH_CM, A4_HEIGHT_CM, DEFAULT_MARGINS
from .utils import calculate_auto_size, split_into_pages

logger = logging.getLogger(__name__)


def set_table_borders(table: Any, visible: bool = False) -> None:
    """Устанавливает видимость границ таблицы."""
    tbl = table._element

    tbl_borders = tbl.find(qn("w:tblBorders"))
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr = tbl.find(qn("w:tblPr"))
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            tbl.insert(0, tbl_pr)
        tbl_pr.insert(0, tbl_borders)

    border_style: str = "single" if visible else "nil"
    border_size: int = 4 if visible else 0
    border_color: str = "C0C0C0" if visible else "auto"

    borders: Dict[str, Dict[str, Any]] = {
        "top": {
            "val": border_style,
            "sz": border_size,
            "space": "0",
            "color": border_color,
        },
        "left": {
            "val": border_style,
            "sz": border_size,
            "space": "0",
            "color": border_color,
        },
        "bottom": {
            "val": border_style,
            "sz": border_size,
            "space": "0",
            "color": border_color,
        },
        "right": {
            "val": border_style,
            "sz": border_size,
            "space": "0",
            "color": border_color,
        },
        "insideH": {
            "val": border_style,
            "sz": border_size,
            "space": "0",
            "color": border_color,
        },
        "insideV": {
            "val": border_style,
            "sz": border_size,
            "space": "0",
            "color": border_color,
        },
    }

    for border_name, border_attrs in borders.items():
        border = tbl_borders.find(qn(f"w:{border_name}"))
        if border is None:
            border = OxmlElement(f"w:{border_name}")
            tbl_borders.append(border)

        for key, value in border_attrs.items():
            border.set(qn(f"w:{key}"), str(value))

    for row in table.rows:
        for cell in row.cells:
            tc_borders = cell._element.find(qn("w:tcBorders"))
            if tc_borders is None:
                tc_borders = OxmlElement("w:tcBorders")
                tc_pr = cell._element.find(qn("w:tcPr"))
                if tc_pr is None:
                    tc_pr = OxmlElement("w:tcPr")
                    cell._element.append(tc_pr)
                tc_pr.append(tc_borders)

            for border_name in ["top", "left", "bottom", "right"]:
                border = tc_borders.find(qn(f"w:{border_name}"))
                if border is None:
                    border = OxmlElement(f"w:{border_name}")
                    tc_borders.append(border)
                border.set(qn("w:val"), "nil")
                border.set(qn("w:sz"), "0")
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), "auto")


def get_image_width_for_table(
    rows: int,
    cols: int,
    image_size_option: str,
    page_width_cm: float = A4_WIDTH_CM,
    page_height_cm: float = A4_HEIGHT_CM,
    margins: Optional[Dict[str, float]] = None,
) -> Cm:
    """Возвращает ширину изображения для таблицы."""
    if margins is None:
        margins = DEFAULT_MARGINS

    if image_size_option != "auto":
        size_cm_tuple: Tuple[float, float] = SIZE_OPTIONS.get(
            image_size_option, (5.0, 5.0)
        )
        return Cm(size_cm_tuple[0])

    usable_page_width: float = page_width_cm - (margins["left"] + margins["right"])
    usable_page_height: float = page_height_cm - (margins["top"] + margins["bottom"])

    size_cm: Tuple[float, float] = calculate_auto_size(
        rows=rows,
        cols=cols,
        page_width_cm=usable_page_width,
        page_height_cm=usable_page_height,
    )
    return Cm(size_cm[0])


def create_single_page_document(
    photos: List[bytes],
    rows: int,
    cols: int,
    table_title: str,
    image_size_option: str = "auto",
) -> Document:
    """Создает одностраничный документ с таблицей из фотографий."""
    doc: Document = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    section = doc.sections[0]
    section.page_width = Cm(A4_WIDTH_CM)
    section.page_height = Cm(A4_HEIGHT_CM)

    section.left_margin = Cm(DEFAULT_MARGINS["left"])
    section.right_margin = Cm(DEFAULT_MARGINS["right"])
    section.top_margin = Cm(DEFAULT_MARGINS["top"])
    section.bottom_margin = Cm(DEFAULT_MARGINS["bottom"])

    if table_title:
        title_paragraph = doc.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = title_paragraph.add_run(table_title)
        run.font.name = "Times New Roman"
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)

        title_paragraph.paragraph_format.space_after = Pt(12)

    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    table.autofit = False
    table.allow_autofit = False

    image_width: Cm = get_image_width_for_table(rows, cols, image_size_option)
    cell_width: Cm = image_width + Cm(0.2)
    cell_height: Cm = image_width + Cm(0.2)

    for row_idx in range(rows):
        for col_idx in range(cols):
            photo_index: int = row_idx * cols + col_idx
            if photo_index < len(photos):
                cell = table.cell(row_idx, col_idx)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

                tc_pr = cell._element.tcPr
                if tc_pr is None:
                    tc_pr = OxmlElement("w:tcPr")
                    cell._element.append(tc_pr)

                tc_mar = OxmlElement("w:tcMar")
                for pos in ["top", "left", "bottom", "right"]:
                    mar = OxmlElement(f"w:{pos}")
                    mar.set(qn("w:w"), "0")
                    mar.set(qn("w:type"), "dxa")
                    tc_mar.append(mar)
                tc_pr.append(tc_mar)

                if photos[photo_index]:
                    paragraph = cell.paragraphs[0]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0

                    run = paragraph.add_run()

                    image_stream: io.BytesIO = io.BytesIO(photos[photo_index])
                    try:
                        run.add_picture(image_stream, width=image_width)
                    except Exception as e:
                        logger.error(f"Ошибка при добавлении изображения: {e}")
                        run.add_text("[Изображение]")
            else:
                cell = table.cell(row_idx, col_idx)
                cell.text = ""

    for row in table.rows:
        row.height_rule = 1
        row.height = cell_height

        for cell in row.cells:
            cell.width = cell_width

    for col_idx in range(cols):
        for row_idx in range(rows):
            try:
                cell = table.cell(row_idx, col_idx)
                cell.width = cell_width
            except Exception:
                pass

    set_table_borders(table, visible=False)

    return doc


def create_multi_page_document(
    photos: List[bytes],
    rows: int,
    cols: int,
    table_title: str,
    image_size_option: str = "auto",
) -> Document:
    """Создает многостраничный документ с таблицами из фотографий."""
    pages: List[List[bytes]] = split_into_pages(photos, rows, cols)

    if len(pages) == 1:
        return create_single_page_document(
            photos, rows, cols, table_title, image_size_option
        )

    first_page: Document = create_single_page_document(
        pages[0],
        rows,
        cols,
        f"{table_title} (стр. 1 из {len(pages)})" if table_title else "",
        image_size_option,
    )

    for page_num, page_photos in enumerate(pages[1:], 2):
        first_page.add_page_break()

        page_doc: Document = create_single_page_document(
            page_photos,
            rows,
            cols,
            f"{table_title} (стр. {page_num} из {len(pages)})"
            if table_title
            else f"Страница {page_num} из {len(pages)}",
            image_size_option,
        )

        for element in page_doc.element.body:
            first_page.element.body.append(element)

    return first_page


create_document_with_table = create_single_page_document
